from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any, Mapping, Sequence

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from psycopg2.extras import RealDictCursor

from gonggong.config.database import get_database_connection


MAX_MARKDOWN_BYTES = 2 * 1024 * 1024
STATION_CODE_PATTERN = re.compile(r"(?<!\d)(\d{12})(?!\d)")


def build_code_convert_filename(original_filename: str | None = None, now: datetime | None = None) -> str:
    timestamp = (now or datetime.now()).strftime("%Y%m%d%H%M%S")
    stem = "赌博分析报告"
    if original_filename:
        cleaned = re.sub(r"[\\/:*?\"<>|]+", "", original_filename.rsplit(".", 1)[0]).strip()
        if cleaned:
            stem = cleaned[:80]
    return f"{stem}_派出所名称转换{timestamp}.docx"


def convert_markdown_station_codes_to_docx(file_bytes: bytes, original_filename: str | None = None) -> io.BytesIO:
    markdown = _decode_markdown(file_bytes)
    codes = extract_station_codes(markdown)
    station_map = fetch_station_name_map(codes)
    converted_markdown = replace_station_codes(markdown, station_map)
    missing_codes = sorted(code for code in codes if code not in station_map)
    return markdown_to_docx(converted_markdown, missing_codes=missing_codes, original_filename=original_filename)


def extract_station_codes(text: str) -> list[str]:
    return sorted(set(STATION_CODE_PATTERN.findall(text or "")))


def fetch_station_name_map(codes: Sequence[str]) -> dict[str, str]:
    clean_codes = [str(code).strip() for code in codes if str(code).strip()]
    if not clean_codes:
        return {}

    placeholders = ", ".join(["%s"] * len(clean_codes))
    sql = f"""
        SELECT
            BTRIM(sspcsdm) AS code,
            COALESCE(MAX(NULLIF(BTRIM(sspcs), '')), BTRIM(sspcsdm)) AS name
        FROM stdata.b_dic_zzjgdm
        WHERE BTRIM(sspcsdm) IN ({placeholders})
        GROUP BY BTRIM(sspcsdm)
    """
    connection = get_database_connection()
    try:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute(sql, clean_codes)
            rows = cursor.fetchall()
        return {
            str(row.get("code") or "").strip(): str(row.get("name") or "").strip()
            for row in rows
            if str(row.get("code") or "").strip()
        }
    finally:
        connection.close()


def replace_station_codes(text: str, station_map: Mapping[str, str]) -> str:
    if not station_map:
        return text

    def _replace(match: re.Match[str]) -> str:
        code = match.group(1)
        return str(station_map.get(code) or code)

    return STATION_CODE_PATTERN.sub(_replace, text)


def markdown_to_docx(markdown: str, *, missing_codes: Sequence[str] | None = None, original_filename: str | None = None) -> io.BytesIO:
    document = Document()
    _configure_document_style(document)

    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line.strip():
            idx += 1
            continue
        if _looks_like_markdown_table(lines, idx):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            _add_markdown_table(document, table_lines)
            continue
        _add_markdown_line(document, line)
        idx += 1

    if missing_codes:
        document.add_page_break()
        document.add_heading("未转换派出所代码", level=1)
        document.add_paragraph("以下代码未在 stdata.b_dic_zzjgdm(sspcsdm, sspcs) 中找到映射，已在正文中保留原代码：")
        for code in missing_codes:
            document.add_paragraph(str(code), style="List Bullet")

    if original_filename:
        section = document.sections[0]
        footer = section.footer.paragraphs[0]
        footer.text = f"来源文件：{original_filename}"

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def _decode_markdown(file_bytes: bytes) -> str:
    if not file_bytes:
        raise ValueError("请上传 markdown 文件")
    if len(file_bytes) > MAX_MARKDOWN_BYTES:
        raise ValueError("markdown 文件不能超过 2MB")

    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _configure_document_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "仿宋"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    normal.font.size = Pt(12)

    for style_name, font_name in (("Heading 1", "黑体"), ("Heading 2", "黑体"), ("Heading 3", "黑体")):
        style = document.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_markdown_line(document: Document, line: str) -> None:
    stripped = line.strip()
    heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 3)
        document.add_heading(_strip_inline_markdown(heading.group(2)), level=level)
        return
    if re.match(r"^[一二三四五六七八九十]+、", stripped):
        document.add_heading(_strip_inline_markdown(stripped), level=1)
        return
    if re.match(r"^（[一二三四五六七八九十]+）", stripped):
        document.add_heading(_strip_inline_markdown(stripped), level=2)
        return
    if re.match(r"^\d+[、.]\s*", stripped):
        document.add_paragraph(_strip_inline_markdown(stripped), style=None)
        return
    if stripped.startswith(("- ", "* ")):
        document.add_paragraph(_strip_inline_markdown(stripped[2:].strip()), style="List Bullet")
        return
    document.add_paragraph(_strip_inline_markdown(stripped))


def _looks_like_markdown_table(lines: Sequence[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    current = lines[index].strip()
    next_line = lines[index + 1].strip()
    return current.startswith("|") and "|" in current[1:] and bool(re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", next_line))


def _add_markdown_table(document: Document, table_lines: Sequence[str]) -> None:
    rows = [_split_markdown_table_row(line) for line in table_lines if line.strip()]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in rows[1]):
        rows.pop(1)
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(max_cols):
            cell = table.cell(row_idx, col_idx)
            cell.text = _strip_inline_markdown(row[col_idx]) if col_idx < len(row) else ""
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _split_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _strip_inline_markdown(text: Any) -> str:
    value = str(text or "")
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"\*([^*]+)\*", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
    return value
