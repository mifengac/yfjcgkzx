from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table


@dataclass(frozen=True)
class ReplaceStats:
    paragraphs_touched: int
    replacements: Dict[str, int]


def _iter_paragraphs_in_tables(tables: Iterable[Table]) -> Iterable[Paragraph]:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_in_cell(cell)


def _iter_paragraphs_in_cell(cell: _Cell) -> Iterable[Paragraph]:
    for paragraph in cell.paragraphs:
        yield paragraph
    yield from _iter_paragraphs_in_tables(cell.tables)


def _iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for paragraph in doc.paragraphs:
        yield paragraph
    yield from _iter_paragraphs_in_tables(doc.tables)


def _replace_in_paragraph(paragraph: Paragraph, mapping: Dict[str, str]) -> Tuple[bool, Dict[str, int]]:
    old_text = paragraph.text
    if not old_text:
        return False, {}

    new_text = old_text
    per_key_counts: Dict[str, int] = {}
    for old, new in mapping.items():
        count = new_text.count(old)
        if count:
            per_key_counts[old] = count
            new_text = new_text.replace(old, new)

    if new_text == old_text:
        return False, {}

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return True, per_key_counts

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True, per_key_counts


def replace_placeholders(docx_path: Path) -> ReplaceStats:
    mapping = {
        "{result1}": "{{result1}}",
        "{result2}": "{{result2}}",
        "{result3}": "{{result3}}",
        "{result4}": "{{result4}}",
        "{result5}": "{{result5}}",
        "{result6}": "{{result6}}",
    }

    doc = Document(str(docx_path))

    paragraphs_touched = 0
    replacements: Dict[str, int] = {k: 0 for k in mapping.keys()}

    for paragraph in _iter_all_paragraphs(doc):
        touched, per_key_counts = _replace_in_paragraph(paragraph, mapping)
        if touched:
            paragraphs_touched += 1
            for key, cnt in per_key_counts.items():
                replacements[key] += cnt

    doc.save(str(docx_path))
    return ReplaceStats(paragraphs_touched=paragraphs_touched, replacements=replacements)


def main() -> None:
    docx_path = Path("jingqing_anjian/templates/biaochezhajie_ribao.docx")
    if not docx_path.exists():
        raise SystemExit(f"模板不存在: {docx_path}")

    stats = replace_placeholders(docx_path)
    print("OK")
    print(f"docx: {docx_path}")
    print(f"paragraphs_touched: {stats.paragraphs_touched}")
    for key, cnt in stats.replacements.items():
        print(f"{key}: {cnt}")


if __name__ == "__main__":
    main()

