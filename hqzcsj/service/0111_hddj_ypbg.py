#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
黄/赌/打架斗殴研判报告（日报）生成脚本

核心变更（按你的最新要求）：
1) 固定映射：yellow=涉黄、gamble=赌博、fight=打架斗殴
2) 同比/环比：上升用“↑xx%”（红色 stat-up），下降用“↓xx%”（绿色 stat-down），不变/0 用“持平”（绿色 stat-down）
3) 删除 SQL/数据库连接：全部数据来源于内网 dsjfx 接口
   - 统计（stats_mock）：来自 /dsjfx/srr/list，取 rows 中 name=="合计" 的行
   - 明细（原 SQL df）：来自 /dsjfx/case/list（字段结构一致）
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
from io import BytesIO
import urllib.error
import urllib.parse
import urllib.request
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from http.cookiejar import CookieJar
from typing import Any, Dict, Iterable, List, Optional, Tuple


from gonggong.config.upstream_province_jingqing import UPSTREAM_PROVINCE_JINGQING_CONFIG

# 从公共配置模块读取省厅警情系统地址
DEFAULT_ORIGIN = UPSTREAM_PROVINCE_JINGQING_CONFIG.get("base_url", "http://68.29.179.170")
DEFAULT_LOGIN_URL = f"{DEFAULT_ORIGIN}{UPSTREAM_PROVINCE_JINGQING_CONFIG.get('login_path', '/dsjfxxb/login')}"
DEFAULT_SRR_URL = f"{DEFAULT_ORIGIN}/dsjfxxb/srr/list"
DEFAULT_CASE_URL = f"{DEFAULT_ORIGIN}/dsjfxxb/case/list"

DEFAULT_USERNAME = UPSTREAM_PROVINCE_JINGQING_CONFIG.get("username", "270378")
DEFAULT_PASSWORD = UPSTREAM_PROVINCE_JINGQING_CONFIG.get("password", "jpx8hLPMyV7EDVX1p9d89Q==")

DEFAULT_LLM_BASE_URL = "http://localhost:8080/v1"
DEFAULT_LLM_MODEL = "police-qwen3-8b"
DEFAULT_TEMPLATE_PATH = os.path.normpath(os.path.join(os.path.dirname(__file__), "report_template.html"))

MAX_ANALYSIS_CHARS = 6000
MOBILE_RE = re.compile(r"^1[3-9]\d{9}$")
EARTH_RADIUS_M = 6371000.0


def html_escape(s: Any) -> str:
    t = "" if s is None else str(s)
    return (
        t.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


@dataclass(frozen=True)
class Category:
    key: str  # yellow / gamble / fight
    title: str
    codes: List[str]


def log(msg: str, *, quiet: bool = False) -> None:
    if quiet:
        return
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{ts}] {msg}")


def _first(items: Iterable[str], default: str = "") -> str:
    for it in items:
        if it:
            return it
    return default


def _parse_dt(s: str) -> datetime:
    s = (s or "").strip()
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            dt = datetime.strptime(s, fmt)
            if fmt == "%Y-%m-%d":
                return datetime(dt.year, dt.month, dt.day, 0, 0, 0)
            return dt
        except ValueError:
            continue
    raise ValueError(f"无法解析时间：{s}，期望格式如 2026-01-10 或 2026-01-10 23:59:59")


def _fmt_dt(dt: datetime) -> str:
    return dt.strftime("%Y-%m-%d %H:%M:%S")


def _date_shift_year(d: date, years: int) -> date:
    try:
        return d.replace(year=d.year + years)
    except ValueError:
        if d.month == 2 and d.day == 29:
            return date(d.year + years, 2, 28)
        raise


def compute_y2y_window(start_dt: datetime, end_dt: datetime) -> Tuple[datetime, datetime]:
    s = start_dt.date()
    e = end_dt.date()
    s2 = _date_shift_year(s, -1)
    e2 = _date_shift_year(e, -1)
    return (
        datetime(s2.year, s2.month, s2.day, 0, 0, 0),
        datetime(e2.year, e2.month, e2.day, 23, 59, 59),
    )


def compute_m2m_window(start_dt: datetime, end_dt: datetime) -> Tuple[datetime, datetime]:
    days = (end_dt.date() - start_dt.date()).days + 1
    m2m_end_date = start_dt.date() - timedelta(days=1)
    m2m_start_date = m2m_end_date - timedelta(days=days - 1)
    return (
        datetime(m2m_start_date.year, m2m_start_date.month, m2m_start_date.day, 0, 0, 0),
        datetime(m2m_end_date.year, m2m_end_date.month, m2m_end_date.day, 23, 59, 59),
    )


def _decode_body(body: bytes, content_type: str = "") -> str:
    m = re.search(r"charset=([a-zA-Z0-9._-]+)", content_type or "")
    candidates = [m.group(1)] if m else []
    candidates += ["utf-8", "gb18030"]
    encoding = _first([c for c in candidates if c], "utf-8")
    for enc in [encoding, "utf-8", "gb18030", "latin-1"]:
        try:
            return body.decode(enc)
        except UnicodeDecodeError:
            continue
    return body.decode("utf-8", errors="replace")


def _request(
    opener: urllib.request.OpenerDirector,
    method: str,
    url: str,
    *,
    headers: Optional[Dict[str, str]] = None,
    data: Optional[bytes] = None,
    timeout: int = 30,
) -> Tuple[int, Dict[str, str], bytes]:
    req = urllib.request.Request(url, data=data, method=method.upper())
    for k, v in (headers or {}).items():
        req.add_header(k, v)
    try:
        with opener.open(req, timeout=timeout) as resp:
            status = getattr(resp, "status", resp.getcode())
            resp_headers = {k: v for k, v in resp.headers.items()}
            body = resp.read()
            return int(status), resp_headers, body
    except urllib.error.HTTPError as e:
        body = e.read() if hasattr(e, "read") else b""
        resp_headers = {k: v for k, v in getattr(e, "headers", {}).items()}
        return int(getattr(e, "code", 0) or 0), resp_headers, body


def _looks_not_logged_in(data: Any) -> bool:
    if not isinstance(data, dict):
        return False
    msg = str(data.get("msg", "") or "")
    return ("未登录" in msg) or ("登录超时" in msg)


def _build_opener() -> urllib.request.OpenerDirector:
    cookie_jar = CookieJar()
    return urllib.request.build_opener(urllib.request.HTTPCookieProcessor(cookie_jar))


def login(opener: urllib.request.OpenerDirector, login_url: str, username: str, password: str, timeout: int) -> None:
    common_headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }
    status, _, _ = _request(opener, "GET", login_url, headers=common_headers, timeout=timeout)
    if status >= 400:
        raise RuntimeError(f"GET 登录页失败：HTTP {status}")

    payload = {"username": username, "password": password, "rememberMe": "true"}
    form = urllib.parse.urlencode(payload).encode("utf-8")
    post_headers = {
        **common_headers,
        "Content-Type": "application/x-www-form-urlencoded",
        "Referer": login_url,
    }
    status, resp_headers, body = _request(opener, "POST", login_url, headers=post_headers, data=form, timeout=timeout)
    if status >= 400:
        preview = _decode_body(body, resp_headers.get("Content-Type", ""))[:200]
        raise RuntimeError(f"POST 登录失败：HTTP {status}, body[:200]={preview}")


def fetch_json_post(
    opener: urllib.request.OpenerDirector,
    url: str,
    payload: Dict[str, str],
    *,
    referer: str,
    origin: str = DEFAULT_ORIGIN,
    timeout: int = 30,
) -> Any:
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
        ),
        "Accept": "application/json, text/javascript, */*; q=0.01",
        "X-Requested-With": "XMLHttpRequest",
        "Content-Type": "application/x-www-form-urlencoded",
        "Origin": origin,
        "Referer": referer,
    }
    form = urllib.parse.urlencode(payload).encode("utf-8")
    status, resp_headers, body = _request(opener, "POST", url, headers=headers, data=form, timeout=timeout)
    text = _decode_body(body, resp_headers.get("Content-Type", ""))
    if status != 200:
        raise RuntimeError(f"请求失败：{url} HTTP {status}\n{text[:500]}")
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        raise RuntimeError(f"响应不是JSON：{url}\n{text[:500]}")


def fetch_srr_rows(
    opener: urllib.request.OpenerDirector,
    srr_url: str,
    *,
    start_dt: datetime,
    end_dt: datetime,
    y2y_start: datetime,
    y2y_end: datetime,
    m2m_start: datetime,
    m2m_end: datetime,
    chara_no_csv: str,
    timeout: int,
) -> List[Dict[str, Any]]:
    payload = {
        "params[startTime]": _fmt_dt(start_dt),
        "params[endTime]": _fmt_dt(end_dt),
        "groupField": "duty_dept_no",
        "caseLevel": "",
        "charaNo": chara_no_csv,
        "chara": "",
        "charaType": "chara_ori",
        "charaLevel": "1",
        "params[y2yStartTime]": _fmt_dt(y2y_start),
        "params[y2yEndTime]": _fmt_dt(y2y_end),
        "dutyDeptNo": "",
        "dutyDeptName": "全部",
        "newRecvType": "",
        "newRecvTypeName": "全部",
        "newCaseSourceNo": "",
        "newCaseSource": "全部",
        "params[m2mStartTime]": _fmt_dt(m2m_start),
        "params[m2mEndTime]": _fmt_dt(m2m_end),
        "params[searchAnd]": "",
        "params[searchOr]": "",
        "params[searchNot]": "",
        "caseContents": "on",
        "replies": "on",
        "pageNum": "NaN",
        "orderByColumn": "",
        "isAsc": "asc",
    }
    data = fetch_json_post(opener, srr_url, payload, referer=f"{DEFAULT_ORIGIN}/dsjfxxb/srr", timeout=timeout)
    if _looks_not_logged_in(data):
        raise RuntimeError("srr/list 返回未登录或超时，请检查登录脚本/账号状态。")
    if not isinstance(data, dict):
        raise RuntimeError("srr/list 响应格式异常（非对象）")
    rows = data.get("rows", []) or []
    if not isinstance(rows, list):
        raise RuntimeError("srr/list rows 非数组")
    return [r for r in rows if isinstance(r, dict)]


def fetch_case_pages(
    opener: urllib.request.OpenerDirector,
    case_url: str,
    *,
    begin_dt: datetime,
    end_dt: datetime,
    new_ori_chara_codes_csv: str,
    timeout: int,
    page_size: int = 200,
    caller_phone_csv: str = "",
) -> List[Dict[str, Any]]:
    rows_all: List[Dict[str, Any]] = []
    page_num = 1
    total: Optional[int] = None

    while True:
        payload = {
            "params[colArray]": "",
            "beginDate": _fmt_dt(begin_dt),
            "endDate": _fmt_dt(end_dt),
            "newCaseSourceNo": "",
            "newCaseSource": "全部",
            "dutyDeptNo": "",
            "dutyDeptName": "全部",
            "newCharaSubclassNo": "",
            "newCharaSubclass": "全部",
            "newOriCharaSubclassNo": new_ori_chara_codes_csv,
            "newOriCharaSubclass": "",
            "caseNo": "",
            "callerName": "",
            "callerPhone": caller_phone_csv,
            "phoneAddress": "",
            "callerIdentity": "",
            "operatorNo": "",
            "operatorName": "",
            "params[isInvalidCase]": "",
            "occurAddress": "",
            "caseMarkNo": "",
            "caseMark": "全部",
            "params[repetitionCase]": "",
            "params[originalDuplicateCase]": "",
            "params[startTimePeriod]": "",
            "params[endTimePeriod]": "",
            "caseContents": "",
            "replies": "",
            "params[sinceRecord]": "",
            "dossierResult": "",
            "params[isVideo]": "",
            "params[isConversation]": "",
            "pageSize": str(page_size),
            "pageNum": str(page_num),
            "orderByColumn": "callTime",
            "isAsc": "desc",
        }
        data = fetch_json_post(opener, case_url, payload, referer=f"{DEFAULT_ORIGIN}/dsjfxxb/case", timeout=timeout)
        if _looks_not_logged_in(data):
            raise RuntimeError("case/list 返回未登录或超时，请检查登录脚本/账号状态。")
        if not isinstance(data, dict):
            raise RuntimeError("case/list 响应格式异常（非对象）")

        if total is None:
            try:
                total = int(data.get("total", 0) or 0)
            except Exception:
                total = 0

        rows = data.get("rows", []) or []
        if not isinstance(rows, list):
            raise RuntimeError("case/list rows 非数组")
        for r in rows:
            if isinstance(r, dict):
                rows_all.append(r)

        if len(rows) < page_size:
            break
        if total is not None and len(rows_all) >= total:
            break

        page_num += 1
        if page_num > 2000:
            break

    return rows_all


def normalize_cmd_name(cmd_name: Any) -> str:
    s = "" if cmd_name is None else str(cmd_name)
    for k in ["云城", "云安", "新兴", "郁南", "罗定"]:
        if k in s:
            return k
    if "云浮市" in s:
        return "市局"
    return s.strip() or "其他"


def is_mobile_phone(s: Any) -> bool:
    if s is None:
        return False
    t = str(s).strip()
    if t in ("", "0000"):
        return False
    return MOBILE_RE.match(t) is not None


def _parse_float(v: Any) -> Optional[float]:
    if v is None:
        return None
    s = str(v).strip()
    if not s:
        return None
    try:
        f = float(s)
    except Exception:
        return None
    if abs(f) < 1e-12:
        return None
    return f


def haversine_m(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    import math

    phi1 = math.radians(lat1)
    phi2 = math.radians(lat2)
    dphi = math.radians(lat2 - lat1)
    dlambda = math.radians(lon2 - lon1)
    a = math.sin(dphi / 2) ** 2 + math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2) ** 2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    return EARTH_RADIUS_M * c


def _extract_line_from(text: str, start_pos: int) -> str:
    end_pos = text.find("\n", start_pos)
    if end_pos == -1:
        return text[start_pos:].strip()
    return text[start_pos:end_pos].strip()


def clean_replies(replies: Any) -> str:
    if replies is None:
        return "未反馈"
    text = str(replies)
    if text.strip() == "":
        return "未反馈"

    pos_feedback = text.rfind("【结警反馈】")
    pos_modify = text.rfind("【结警修改】")
    if pos_feedback != -1 or pos_modify != -1:
        start_pos = max(pos_feedback, pos_modify)
        return _extract_line_from(text, start_pos)

    fallback_patterns = [
        "【过程反馈】",
        "提交处理情况",
        "申请【升级改派】",
        "选择无需出警，无需出警的原因：",
        "选择不出警，不出警原因：",
    ]
    for pattern in fallback_patterns:
        pos = text.rfind(pattern)
        if pos != -1:
            return _extract_line_from(text, pos)

    return "未反馈"


def _extract_total_row(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    for r in rows:
        if str(r.get("name", "") or "") == "合计":
            return r
    return {}


def _parse_int_like(v: Any) -> int:
    s = "" if v is None else str(v)
    m = re.search(r"(\d+)", s)
    if not m:
        return 0
    try:
        return int(m.group(1))
    except Exception:
        return 0


def _format_delta(v: Any) -> str:
    s = "" if v is None else str(v).strip()
    if not s:
        return "持平"
    if "持平" in s:
        return "持平"

    m = re.search(r"([+-]?)\s*(\d+(?:\.\d+)?)\s*%?", s)
    if not m:
        return s
    sign = m.group(1) or ""
    num = m.group(2) or "0"
    try:
        f = float(num)
    except Exception:
        return s

    if abs(f) < 1e-9:
        return "持平"
    if sign == "-":
        return f"↓{num}%"
    return f"↑{num}%"


def build_stats_mock_from_srr_total(total_row: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "count": _parse_int_like(total_row.get("presentCycle", "")),
        "yoy": _format_delta(total_row.get("y2yProportion", "")),
        "mom": _format_delta(total_row.get("m2mProportion", "")),
    }


def summarize_one_case(llm: "LLMClient", case_contents: str, replies: str) -> str:
    sys_p = "你是一名公安情报员。请将报警内容与处警情况压缩成一句话摘要（不超过40字），不需要加序号。"
    prompt = f"报警内容：{case_contents}\n处警情况：{replies}"
    return llm.chat(prompt, system_prompt=sys_p, max_tokens=128).strip()


def make_case_summary(llm: Optional["LLMClient"], case_contents: Any, replies: Any) -> str:
    contents = "" if case_contents is None else str(case_contents)
    if llm is None:
        return contents.strip()
    rep = clean_replies(replies)
    try:
        return summarize_one_case(llm, contents, rep)
    except Exception:
        return contents.strip()


def analyze_problems(llm: "LLMClient", summaries_text: str) -> str:
    sys_p = "你是一名公安情报员。请基于警情摘要，输出3条“重点问题”，用1.2.3.编号。"
    return llm.chat(summaries_text, system_prompt=sys_p, max_tokens=256).strip()


def analyze_measures(llm: "LLMClient", summaries_text: str, problems_text: str) -> str:
    sys_p = "你是一名公安情报员。请结合重点问题与警情摘要，输出3条“下一步措施”，用1.2.3.编号。"
    prompt = f"重点问题：\n{problems_text}\n\n警情摘要：\n{summaries_text}"
    return llm.chat(prompt, system_prompt=sys_p, max_tokens=256).strip()


def compute_basic_stats(cases: List[Dict[str, Any]]) -> Dict[str, Any]:
    total = len(cases)
    regions = []
    seen = set()
    for c in cases:
        r = str(c.get("地区", "") or "").strip()
        if not r or r in seen:
            continue
        seen.add(r)
        regions.append(r)

    parts: List[str] = []
    for region in regions:
        region_cases = [c for c in cases if c.get("地区") == region]
        cnt_yellow = sum(1 for c in region_cases if c.get("警情性质") == "涉黄")
        cnt_gamble = sum(1 for c in region_cases if c.get("警情性质") == "赌博")
        cnt_fight = sum(1 for c in region_cases if c.get("警情性质") == "打架斗殴")
        sub: List[str] = []
        if cnt_yellow:
            sub.append(f"涉黄{cnt_yellow}起")
        if cnt_gamble:
            sub.append(f"赌博{cnt_gamble}起")
        if cnt_fight:
            sub.append(f"打架斗殴{cnt_fight}起")
        if sub:
            parts.append(f"{region}{'，'.join(sub)}")

    type_summary_str = "；".join(parts) + ("。" if parts else "")
    return {"total_count": total, "type_summary_str": type_summary_str}


def build_nested_list(cases: List[Dict[str, Any]], type_value: str) -> List[Dict[str, Any]]:
    filtered = [c for c in cases if c.get("警情性质") == type_value]
    if not filtered:
        return []

    regions: Dict[str, Dict[str, List[str]]] = {}
    for c in filtered:
        region = str(c.get("地区", "") or "").strip()
        station = str(c.get("派出所", "") or "").strip()
        summary = str(c.get("case_summary", "") or "").strip()
        if not region:
            region = "其他"
        if not station:
            station = "未知派出所"
        regions.setdefault(region, {}).setdefault(station, []).append(summary)

    out: List[Dict[str, Any]] = []
    for region_name, station_map in regions.items():
        stations = []
        for station_name, summaries in station_map.items():
            summary_text = "；".join(s for s in summaries if s)
            stations.append(
                {
                    "name": station_name,
                    "count": len(summaries),
                    "cases": [s for s in summaries if s],
                    "summary_text": summary_text,
                }
            )
        out.append({"cmdname": region_name, "stations": stations})
    return out


class LLMClient:
    def __init__(
        self,
        base_url: str,
        *,
        model: str,
        timeout: int = 60,
        cache_path: str = "",
    ) -> None:
        self.base_url = (base_url or "").rstrip("/")
        self.model = model
        self.timeout = int(timeout)
        self.cache_path = cache_path
        self._cache: Dict[str, str] = {}
        if self.cache_path:
            self._load_cache()

    def _load_cache(self) -> None:
        if not os.path.exists(self.cache_path):
            return
        try:
            with open(self.cache_path, "r", encoding="utf-8") as f:
                self._cache = json.load(f) or {}
        except Exception:
            self._cache = {}

    def _save_cache(self) -> None:
        if not self.cache_path:
            return
        os.makedirs(os.path.dirname(self.cache_path) or ".", exist_ok=True)
        tmp = self.cache_path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(self._cache, f, ensure_ascii=False, indent=2)
        os.replace(tmp, self.cache_path)

    def chat(self, prompt: str, *, system_prompt: str, max_tokens: int = 256) -> str:
        key = hashlib.sha1((system_prompt + "\n" + prompt).encode("utf-8")).hexdigest()
        if key in self._cache:
            return self._cache[key]

        url = f"{self.base_url}/chat/completions"
        payload = {
            "model": self.model,
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],
            "max_tokens": int(max_tokens),
        }
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        req = urllib.request.Request(url, data=data, method="POST")
        req.add_header("Content-Type", "application/json")

        try:
            with urllib.request.urlopen(req, timeout=self.timeout) as resp:
                body = resp.read()
        except Exception as e:
            raise RuntimeError(f"LLM请求失败：{e}")

        try:
            obj = json.loads(body.decode("utf-8", errors="replace"))
            content = obj["choices"][0]["message"]["content"]
        except Exception as e:
            raise RuntimeError(f"LLM响应解析失败：{e}")

        out = str(content or "").strip()
        self._cache[key] = out
        self._save_cache()
        return out


def render_html(template_path: str, context: Dict[str, Any]) -> str:
    try:
        from jinja2 import Environment, FileSystemLoader  # type: ignore
    except Exception as e:
        raise RuntimeError(f"缺少依赖 jinja2，无法渲染模板：{e}")

    tpl_path = os.path.abspath(template_path)
    loader_dir = os.path.dirname(tpl_path) or "."
    tpl_name = os.path.basename(tpl_path)
    env = Environment(loader=FileSystemLoader(loader_dir), autoescape=False)
    template = env.get_template(tpl_name)
    return template.render(context)


def _text(v: Any) -> str:
    return "" if v is None else str(v)


def _iter_non_empty_lines(text: Any) -> List[str]:
    lines = []
    for line in _text(text).replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        s = line.strip()
        if s:
            lines.append(s)
    return lines


def render_report_docx(context: Dict[str, Any]) -> bytes:
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception as e:
        raise RuntimeError(f"缺少依赖 python-docx，无法导出 DOCX：{e}")

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(11)
    normal_style.font.name = "FangSong"
    try:
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    except Exception:
        pass

    def add_h1(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_h2(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)

    def add_text(text: str, indent: bool = False) -> None:
        p = doc.add_paragraph(text)
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)

    def add_srr_table(title: str, rows: List[Dict[str, Any]]) -> None:
        add_text(title)
        headers = ["序号", "单位", "本期数", "同比值", "同比", "环比值", "环比"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
        if not rows:
            row_cells = table.add_row().cells
            row_cells[0].text = "-"
            row_cells[1].text = "无数据"
            for i in range(2, len(headers)):
                row_cells[i].text = "-"
            return
        for idx, row in enumerate(rows, start=1):
            cells = table.add_row().cells
            cells[0].text = str(idx)
            cells[1].text = _text(row.get("name", ""))
            cells[2].text = str(row.get("presentCycle", 0) or 0)
            cells[3].text = str(row.get("upperY2yCycle", 0) or 0)
            cells[4].text = _text(row.get("y2yProportion", "持平")) or "持平"
            cells[5].text = str(row.get("upperM2mCycle", 0) or 0)
            cells[6].text = _text(row.get("m2mProportion", "持平")) or "持平"
            if _text(row.get("name", "")) == "合计":
                for c in cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.bold = True

    def add_nested_section(title: str, regions: List[Dict[str, Any]]) -> None:
        add_h2(title)
        if not regions:
            add_text("无相关警情。")
            return
        for region in regions:
            add_text(_text(region.get("cmdname", "")))
            for station in region.get("stations", []) or []:
                summary = _text(station.get("summary_text", ""))
                add_text(f"{_text(station.get('name', ''))}（{_text(station.get('count', 0))}）起：{summary}", indent=True)

    def add_repeat_section(title: str, items: List[Dict[str, Any]], empty_text: str) -> None:
        add_h2(title)
        if not items:
            add_text(empty_text)
            return
        for it in items:
            add_text(f"{_text(it.get('phone', ''))}（{_text(it.get('count', 0))}次）：")
            for line in it.get("lines", []) or []:
                add_text(_text(line), indent=True)

    def add_multi_addr_section(title: str, items: List[Dict[str, Any]]) -> None:
        add_h2(title)
        if not items:
            add_text("无多次报警地址。")
            return
        for it in items:
            examples = "；".join(_text(x) for x in (it.get("examples", []) or []))
            add_text(f"{_text(it.get('occurAddress', ''))}（{_text(it.get('count', 0))}次）：{examples}", indent=True)

    add_h1(f"涉黄、赌、打架斗殴警情研判日报 ({_text(context.get('report_date', ''))})")
    add_h2("一、当日警情概况")
    add_text(
        f"全市共接报涉黄、赌、打架斗殴类警情 {_text(context.get('total_count', 0))} 起。"
        f"其中：{_text(context.get('type_summary_str', ''))}",
        indent=True,
    )

    srr_start = _text(context.get("srr_table_start_date", ""))
    srr_end = _text(context.get("srr_table_end_date", ""))
    srr_rows_by_cat = context.get("srr_rows_by_cat", {}) or {}
    add_h2("二、分类警情统计表")
    add_srr_table(f"{srr_start}至{srr_end}打架斗殴警情统计", list(srr_rows_by_cat.get("fight", []) or []))
    add_srr_table(f"{srr_start}至{srr_end}涉黄警情统计", list(srr_rows_by_cat.get("yellow", []) or []))
    add_srr_table(f"{srr_start}至{srr_end}赌博警情统计", list(srr_rows_by_cat.get("gamble", []) or []))

    add_nested_section("三、打架斗殴类警情", list(context.get("fight_list", []) or []))
    add_nested_section("四、赌博类警情", list(context.get("gamble_list", []) or []))
    add_nested_section("五、涉黄类警情", list(context.get("sex_list", []) or []))

    repeat_min_count = _text(context.get("repeat_min_count", 2))
    add_repeat_section(
        f"六、重复报警电话（{repeat_min_count}次及以上）- 打架斗殴",
        list(context.get("repeat_fight", []) or []),
        "无重复报警电话。",
    )
    add_repeat_section(
        f"七、重复报警电话（{repeat_min_count}次及以上）- 赌博",
        list(context.get("repeat_gamble", []) or []),
        "无重复报警电话。",
    )
    add_repeat_section(
        f"八、重复报警电话（{repeat_min_count}次及以上）- 涉黄",
        list(context.get("repeat_yellow", []) or []),
        "无重复报警电话。",
    )

    add_multi_addr_section("九、多次报警地址（近3个月半径50米）- 打架斗殴", list(context.get("multi_addr_fight", []) or []))
    add_multi_addr_section("十、多次报警地址（近3个月半径50米）- 赌博", list(context.get("multi_addr_gamble", []) or []))
    add_multi_addr_section("十一、多次报警地址（近3个月半径50米）- 涉黄", list(context.get("multi_addr_yellow", []) or []))

    add_h2("十二、重点问题")
    probs = _iter_non_empty_lines(context.get("llm_analysis_problems", ""))
    if probs:
        for line in probs:
            add_text(line, indent=True)
    else:
        add_text("")

    add_h2("十三、下一步措施")
    measures = _iter_non_empty_lines(context.get("llm_analysis_measures", ""))
    if measures:
        for line in measures:
            add_text(line, indent=True)
    else:
        add_text("")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def render_report_pdf(context: Dict[str, Any]) -> bytes:
    try:
        from reportlab.lib import colors  # type: ignore
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle  # type: ignore
    except Exception as e:
        raise RuntimeError(f"缺少依赖 reportlab，无法导出 PDF：{e}")

    font_name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        font_name = "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "RptTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        leading=24,
        alignment=1,
    )
    h_style = ParagraphStyle(
        "RptHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=13,
        leading=18,
    )
    n_style = ParagraphStyle(
        "RptNormal",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10.5,
        leading=16,
    )

    def para(text: Any, style: ParagraphStyle) -> Paragraph:
        safe = html_escape(_text(text)).replace("\n", "<br/>")
        return Paragraph(safe, style)

    def make_srr_table(title: str, rows: List[Dict[str, Any]]) -> List[Any]:
        out: List[Any] = [para(title, h_style), Spacer(1, 4)]
        data = [["序号", "单位", "本期数", "同比值", "同比", "环比值", "环比"]]
        if not rows:
            data.append(["-", "无数据", "-", "-", "-", "-", "-"])
        else:
            for idx, row in enumerate(rows, start=1):
                data.append(
                    [
                        str(idx),
                        _text(row.get("name", "")),
                        str(row.get("presentCycle", 0) or 0),
                        str(row.get("upperY2yCycle", 0) or 0),
                        _text(row.get("y2yProportion", "持平")) or "持平",
                        str(row.get("upperM2mCycle", 0) or 0),
                        _text(row.get("m2mProportion", "持平")) or "持平",
                    ]
                )
        table = Table(data, repeatRows=1, colWidths=[28, 140, 58, 58, 58, 58, 58])
        style = TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), font_name, 9.5),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f8fafc")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
        for i, row in enumerate(rows, start=1):
            if _text(row.get("name", "")) == "合计":
                style.add("BACKGROUND", (0, i), (-1, i), colors.HexColor("#fff7ed"))
                style.add("FONTNAME", (0, i), (-1, i), font_name)
        table.setStyle(style)
        out.append(table)
        out.append(Spacer(1, 10))
        return out

    story: List[Any] = []
    story.append(para(f"涉黄、赌、打架斗殴警情研判日报 ({_text(context.get('report_date', ''))})", title_style))
    story.append(Spacer(1, 10))

    story.append(para("一、当日警情概况", h_style))
    story.append(
        para(
            f"全市共接报涉黄、赌、打架斗殴类警情 {_text(context.get('total_count', 0))} 起。"
            f"其中：{_text(context.get('type_summary_str', ''))}",
            n_style,
        )
    )
    story.append(Spacer(1, 8))

    srr_start = _text(context.get("srr_table_start_date", ""))
    srr_end = _text(context.get("srr_table_end_date", ""))
    srr_rows_by_cat = context.get("srr_rows_by_cat", {}) or {}
    story.append(para("二、分类警情统计表", h_style))
    story.extend(make_srr_table(f"{srr_start}至{srr_end}打架斗殴警情统计", list(srr_rows_by_cat.get("fight", []) or [])))
    story.extend(make_srr_table(f"{srr_start}至{srr_end}涉黄警情统计", list(srr_rows_by_cat.get("yellow", []) or [])))
    story.extend(make_srr_table(f"{srr_start}至{srr_end}赌博警情统计", list(srr_rows_by_cat.get("gamble", []) or [])))

    def add_nested(title: str, regions: List[Dict[str, Any]]) -> None:
        story.append(para(title, h_style))
        if not regions:
            story.append(para("无相关警情。", n_style))
            return
        for region in regions:
            story.append(para(_text(region.get("cmdname", "")), n_style))
            for station in region.get("stations", []) or []:
                story.append(
                    para(
                        f"{_text(station.get('name', ''))}（{_text(station.get('count', 0))}）起：{_text(station.get('summary_text', ''))}",
                        n_style,
                    )
                )

    add_nested("三、打架斗殴类警情", list(context.get("fight_list", []) or []))
    add_nested("四、赌博类警情", list(context.get("gamble_list", []) or []))
    add_nested("五、涉黄类警情", list(context.get("sex_list", []) or []))

    def add_repeat(title: str, items: List[Dict[str, Any]], empty_text: str) -> None:
        story.append(para(title, h_style))
        if not items:
            story.append(para(empty_text, n_style))
            return
        for it in items:
            story.append(para(f"{_text(it.get('phone', ''))}（{_text(it.get('count', 0))}次）：", n_style))
            for line in it.get("lines", []) or []:
                story.append(para(_text(line), n_style))

    min_cnt = _text(context.get("repeat_min_count", 2))
    add_repeat(f"六、重复报警电话（{min_cnt}次及以上）- 打架斗殴", list(context.get("repeat_fight", []) or []), "无重复报警电话。")
    add_repeat(f"七、重复报警电话（{min_cnt}次及以上）- 赌博", list(context.get("repeat_gamble", []) or []), "无重复报警电话。")
    add_repeat(f"八、重复报警电话（{min_cnt}次及以上）- 涉黄", list(context.get("repeat_yellow", []) or []), "无重复报警电话。")

    def add_multi(title: str, items: List[Dict[str, Any]]) -> None:
        story.append(para(title, h_style))
        if not items:
            story.append(para("无多次报警地址。", n_style))
            return
        for it in items:
            examples = "；".join(_text(x) for x in (it.get("examples", []) or []))
            story.append(para(f"{_text(it.get('occurAddress', ''))}（{_text(it.get('count', 0))}次）：{examples}", n_style))

    add_multi("九、多次报警地址（近3个月半径50米）- 打架斗殴", list(context.get("multi_addr_fight", []) or []))
    add_multi("十、多次报警地址（近3个月半径50米）- 赌博", list(context.get("multi_addr_gamble", []) or []))
    add_multi("十一、多次报警地址（近3个月半径50米）- 涉黄", list(context.get("multi_addr_yellow", []) or []))

    story.append(para("十二、重点问题", h_style))
    probs = _iter_non_empty_lines(context.get("llm_analysis_problems", ""))
    if probs:
        for line in probs:
            story.append(para(line, n_style))
    else:
        story.append(para("", n_style))
    story.append(para("十三、下一步措施", h_style))
    measures = _iter_non_empty_lines(context.get("llm_analysis_measures", ""))
    if measures:
        for line in measures:
            story.append(para(line, n_style))
    else:
        story.append(para("", n_style))

    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, leftMargin=28, rightMargin=28, topMargin=28, bottomMargin=28)
    doc.build(story)
    bio.seek(0)
    return bio.read()


def build_repeat_phone_list(
    opener: urllib.request.OpenerDirector,
    *,
    case_url: str,
    end_dt: datetime,
    chara_no_csv: str,
    seed_phones: List[str],
    llm: Optional["LLMClient"],
    timeout: int,
    page_size: int,
    min_count: int = 2,
    recent_n: int = 2,
) -> List[Dict[str, Any]]:
    phones = [p for p in seed_phones if is_mobile_phone(p)]
    if not phones:
        return []

    fixed_begin = datetime(2020, 1, 1, 0, 0, 0)
    phone_csv = ",".join(phones)
    rows = fetch_case_pages(
        opener,
        case_url,
        begin_dt=fixed_begin,
        end_dt=end_dt,
        new_ori_chara_codes_csv=chara_no_csv,
        caller_phone_csv=phone_csv,
        timeout=timeout,
        page_size=page_size,
    )

    buckets: Dict[str, List[Dict[str, Any]]] = {}
    for r in rows:
        phone = str(r.get("callerPhone", "") or "").strip()
        if not is_mobile_phone(phone):
            continue
        buckets.setdefault(phone, []).append(r)

    items: List[Dict[str, Any]] = []
    for phone, rs in buckets.items():
        if len(rs) < int(min_count):
            continue
        rs_sorted = sorted(rs, key=lambda x: str(x.get("callTime", "") or ""), reverse=True)
        lines: List[str] = []
        for rr in rs_sorted[: max(1, int(recent_n))]:
            call_time = str(rr.get("callTime", "") or "").strip()
            cmd_name = normalize_cmd_name(rr.get("cmdName", ""))
            duty_dept_name = str(rr.get("dutyDeptName", "") or "").strip()
            summary = make_case_summary(llm, rr.get("caseContents", ""), rr.get("replies", ""))
            line = f"{call_time},{cmd_name},{duty_dept_name}{summary}"
            lines.append(line)

        html_lines = "<br/>".join(f"&emsp;&emsp;{html_escape(x)}" for x in lines if x)
        items.append({"phone": phone, "count": len(rs), "lines": lines, "html": html_lines})

    items.sort(key=lambda x: int(x.get("count", 0)), reverse=True)
    return items


def build_multi_address_items(
    *,
    recent_rows: List[Dict[str, Any]],
    radius_m: float = 50.0,
    min_count: int = 2,
) -> List[Dict[str, Any]]:
    points: List[Tuple[float, float, Dict[str, Any]]] = []
    for r in recent_rows:
        lng = _parse_float(r.get("lngOfLocate"))
        lat = _parse_float(r.get("latOfLocate"))
        if lng is None or lat is None:
            continue
        points.append((lat, lng, r))
    if len(points) < int(min_count):
        return []

    import math

    lat0 = sum(p[0] for p in points) / float(len(points))
    lon_factor = 111320.0 * math.cos(math.radians(lat0))
    lat_factor = 111320.0
    cell = max(1.0, float(radius_m))

    def to_cell(lat: float, lon: float) -> Tuple[int, int]:
        x = lon * lon_factor
        y = lat * lat_factor
        return (int(x // cell), int(y // cell))

    cell_map: Dict[Tuple[int, int], List[int]] = {}
    for idx, (lat, lon, _) in enumerate(points):
        cell_map.setdefault(to_cell(lat, lon), []).append(idx)

    visited = [False] * len(points)
    clusters: List[List[int]] = []

    for i, (lat_i, lon_i, _) in enumerate(points):
        if visited[i]:
            continue
        visited[i] = True
        q = [i]
        comp = [i]
        while q:
            cur = q.pop()
            lat_c, lon_c, _ = points[cur]
            cx, cy = to_cell(lat_c, lon_c)
            for dx in (-1, 0, 1):
                for dy in (-1, 0, 1):
                    cand = cell_map.get((cx + dx, cy + dy), [])
                    for j in cand:
                        if visited[j]:
                            continue
                        lat_j, lon_j, _ = points[j]
                        if haversine_m(lat_c, lon_c, lat_j, lon_j) <= float(radius_m):
                            visited[j] = True
                            q.append(j)
                            comp.append(j)
        clusters.append(comp)

    items: List[Dict[str, Any]] = []
    for comp in clusters:
        if len(comp) < int(min_count):
            continue
        rows = [points[idx][2] for idx in comp]
        rows_sorted = sorted(rows, key=lambda x: str(x.get("callTime", "") or ""), reverse=True)

        addr_counts: Dict[str, int] = {}
        for rr in rows:
            addr = str(rr.get("occurAddress", "") or "").strip()
            if not addr:
                continue
            addr_counts[addr] = addr_counts.get(addr, 0) + 1
        if addr_counts:
            occur_address = max(addr_counts.items(), key=lambda kv: kv[1])[0]
        else:
            occur_address = "未知地址"

        examples: List[str] = []
        for rr in rows_sorted:
            call_time = str(rr.get("callTime", "") or "").strip()
            addr = str(rr.get("occurAddress", "") or "").strip()
            if not call_time or not addr:
                continue
            examples.append(f"{call_time}:{addr}")
            if len(examples) >= 2:
                break

        items.append({"occurAddress": occur_address, "count": len(rows), "examples": examples})

    items.sort(key=lambda x: int(x.get("count", 0)), reverse=True)
    return items


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--start-date", default="2026-02-02")
    parser.add_argument("--end-date", default="2026-02-02")
    parser.add_argument("--format", default="html", help="输出格式：html/docx/pdf")
    parser.add_argument("--login-url", default=DEFAULT_LOGIN_URL)
    parser.add_argument("--srr-url", default=DEFAULT_SRR_URL)
    parser.add_argument("--case-url", default=DEFAULT_CASE_URL)
    parser.add_argument("--username", default=DEFAULT_USERNAME)
    parser.add_argument("--password", default=DEFAULT_PASSWORD)
    parser.add_argument("--timeout", type=int, default=30)
    parser.add_argument("--page-size", type=int, default=200)
    parser.add_argument("--no-llm", action="store_true")
    parser.add_argument("--llm-base-url", default=DEFAULT_LLM_BASE_URL)
    parser.add_argument("--llm-model", default=DEFAULT_LLM_MODEL)
    parser.add_argument("--llm-cache", default="0111_hddj_ypbg_llm_cache.json")
    parser.add_argument("--template", default=DEFAULT_TEMPLATE_PATH)
    parser.add_argument("--output", default="")
    parser.add_argument("--quiet", action="store_true")
    parser.add_argument("--repeat_min_count", type=int, default=2, help="重复报警电话最小次数（默认2=2次及以上）")
    parser.add_argument("--topn_repeat_phone", type=int, default=200, help="参与重复查询的手机号去重后最多取前N个")
    parser.add_argument("--repeat_recent_n", type=int, default=2, help="每个重复手机号仅展示最近N条（默认2）")
    parser.add_argument(
        "--srr-from-year-start",
        dest="srr_from_year_start",
        action="store_true",
        help="仅对 srr/list 生效：将开始时间固定为 end-date 当年的 1 月 1 日 00:00:00（默认开启）",
    )
    parser.add_argument(
        "--no-srr-from-year-start",
        dest="srr_from_year_start",
        action="store_false",
        help="仅对 srr/list 生效：使用 --start-date 作为开始时间",
    )
    parser.set_defaults(srr_from_year_start=True)
    args = parser.parse_args()
    quiet = bool(args.quiet)

    log("步骤：解析时间参数", quiet=quiet)
    start_dt = _parse_dt(args.start_date)
    end_dt = _parse_dt(args.end_date)
    # start_dt = datetime(start_dt.year, start_dt.month, start_dt.day, 0, 0, 0)
    # end_dt = datetime(end_dt.year, end_dt.month, end_dt.day, 23, 59, 59)
    if args.srr_from_year_start:
        srr_start_dt = datetime(end_dt.year, 1, 1, 0, 0, 0)
    else:
        srr_start_dt = start_dt

    report_date_str = start_dt.strftime("%Y年%m月%d日")
    y2y_start, y2y_end = compute_y2y_window(srr_start_dt, end_dt)
    m2m_start, m2m_end = compute_m2m_window(srr_start_dt, end_dt)
    log(f"时间段：{start_dt} ~ {end_dt}", quiet=quiet)
    log(f"srr统计区间：{srr_start_dt} ~ {end_dt}", quiet=quiet)
    log(f"同比区间：{y2y_start} ~ {y2y_end}", quiet=quiet)
    log(f"环比区间：{m2m_start} ~ {m2m_end}", quiet=quiet)

    categories = [
        Category(
            key="fight",
            title="打架斗殴",
            codes=[
                "02010899",
                "02010803",
                "02010802",
                "02010801",
                "01050102",
                "02010800",
                "01030300",
                "02031000",
                "02030100",
            ],
        ),
        Category(
            key="yellow",
            title="涉黄",
            codes=[
                "09020100",
                "09020000",
                "02051899",
                "02051809",
                "02051808",
                "02051807",
                "02051806",
                "02051805",
                "02051804",
                "02051803",
                "02051802",
                "02051801",
                "02051800",
                "01051200",
                "01051199",
                "01051104",
                "01051103",
                "01051102",
                "01051101",
                "01051100",
                "09029900",
                "09020500",
                "09020400",
                "09020300",
                "09020200",
            ],
        ),
        Category(
            key="gamble",
            title="赌博",
            codes=[
                "09019900",
                "09010600",
                "09010500",
                "09010400",
                "09010300",
                "09010200",
                "09010100",
                "09010000",
                "02052099",
                "02052004",
                "02052003",
                "02052002",
                "02052001",
                "02052000",
                "01050499",
                "01050405",
                "01050404",
                "01050403",
                "01050402",
                "01050401",
                "01050400",
            ],
        ),
    ]

    opener = _build_opener()
    log("步骤：登录内网系统", quiet=quiet)
    login(opener, args.login_url, args.username, args.password, args.timeout)
    log("登录请求已提交（后续接口以返回为准校验）", quiet=quiet)

    stats_mock: Dict[str, Any] = {"yellow": {}, "gamble": {}, "fight": {}}
    srr_rows_by_cat: Dict[str, List[Dict[str, Any]]] = {"yellow": [], "gamble": [], "fight": []}
    cases: List[Dict[str, Any]] = []
    seed_phones_by_cat: Dict[str, List[str]] = {c.key: [] for c in categories}

    for cat in categories:
        chara_no_csv = ",".join(cat.codes)
        log(f"步骤：获取统计（srr/list）- {cat.title}", quiet=quiet)
        rows = fetch_srr_rows(
            opener,
            args.srr_url,
            start_dt=srr_start_dt,
            end_dt=end_dt,
            y2y_start=y2y_start,
            y2y_end=y2y_end,
            m2m_start=m2m_start,
            m2m_end=m2m_end,
            chara_no_csv=chara_no_csv,
            timeout=args.timeout,
        )
        srr_rows_by_cat[cat.key] = rows
        total_row = _extract_total_row(rows)
        stats_mock[cat.key] = build_stats_mock_from_srr_total(total_row)

        log(f"步骤：获取明细（case/list）- {cat.title}", quiet=quiet)
        rows_case = fetch_case_pages(
            opener,
            args.case_url,
            begin_dt=start_dt,
            end_dt=end_dt,
            new_ori_chara_codes_csv=chara_no_csv,
            timeout=args.timeout,
            page_size=args.page_size,
        )
        for r in rows_case:
            phone = str(r.get("callerPhone", "") or "").strip()
            if phone:
                seed_phones_by_cat[cat.key].append(phone)
            cases.append(
                {
                    "报警时间": r.get("callTime", ""),
                    "地区": normalize_cmd_name(r.get("cmdName", "")),
                    "派出所": r.get("dutyDeptName", "") or "",
                    "警情地址": r.get("occurAddress", "") or "",
                    "报警内容": r.get("caseContents", "") or "",
                    "处警情况": clean_replies(r.get("replies", "")),
                    "警情性质": cat.title,
                    "callerPhone": phone,
                }
            )

    # LLM：逐条摘要 + 研判分析
    llm: Optional[LLMClient] = None
    if not args.no_llm:
        log(f"步骤：初始化本地大模型客户端 base_url={args.llm_base_url}", quiet=quiet)
        llm = LLMClient(args.llm_base_url, model=args.llm_model, timeout=60, cache_path=args.llm_cache)

    if llm is None or not cases:
        log("步骤：跳过大模型（--no-llm 或无数据），使用规则拼接摘要/默认研判", quiet=quiet)
        for c in cases:
            c["case_summary"] = f"{c.get('报警内容','')} (处理:{c.get('处警情况','')})".strip()
        llm_probs = "1. 重点区域存在治安风险点。\n2. 涉赌警情有聚集趋势。\n3. 深夜打架斗殴警情仍需关注。"
        llm_measures = "1. 加强巡逻防控。\n2. 开展专项打击。\n3. 强化宣传教育。"
    else:
        try:
            log("步骤：逐条摘要（LLM）", quiet=quiet)
            for c in cases:
                c["case_summary"] = summarize_one_case(
                    llm, str(c.get("报警内容", "")), str(c.get("处警情况", ""))
                )
            all_summaries = "\n".join(str(c.get("case_summary", "")) for c in cases)[:MAX_ANALYSIS_CHARS]
            log("步骤：生成研判分析（LLM）- 重点问题/下一步措施", quiet=quiet)
            llm_probs = analyze_problems(llm, all_summaries)
            llm_measures = analyze_measures(llm, all_summaries, llm_probs)
        except Exception as exc:
            log(f"LLM不可用，降级为规则摘要：{exc}", quiet=quiet)
            for c in cases:
                c["case_summary"] = str(c.get("报警内容", "") or "").strip()
            llm_probs = ""
            llm_measures = ""
            llm = None

    log("步骤：重复报警电话（2次及以上）统计", quiet=quiet)
    repeat_by_cat: Dict[str, List[Dict[str, Any]]] = {"yellow": [], "gamble": [], "fight": []}
    for cat in categories:
        raw_phones = seed_phones_by_cat.get(cat.key, [])
        # 去重 + 过滤非手机号
        seen = set()
        phones: List[str] = []
        for p in raw_phones:
            p = str(p or "").strip()
            if not p or p in seen:
                continue
            seen.add(p)
            if not is_mobile_phone(p):
                continue
            phones.append(p)
            if args.topn_repeat_phone > 0 and len(phones) >= int(args.topn_repeat_phone):
                break

        if not phones:
            repeat_by_cat[cat.key] = []
            continue

        chara_no_csv = ",".join(cat.codes)
        repeat_by_cat[cat.key] = build_repeat_phone_list(
            opener,
            case_url=args.case_url,
            end_dt=end_dt,
            chara_no_csv=chara_no_csv,
            seed_phones=phones,
            llm=llm,
            timeout=args.timeout,
            page_size=args.page_size,
            min_count=int(args.repeat_min_count),
            recent_n=int(args.repeat_recent_n),
        )

    log("步骤：多次报警地址（近3个月半径50米聚类，2次及以上）统计", quiet=quiet)
    recent_begin_dt = end_dt - timedelta(days=90)
    recent_begin_dt = datetime(recent_begin_dt.year, recent_begin_dt.month, recent_begin_dt.day, 0, 0, 0)
    multi_addr_by_cat: Dict[str, List[Dict[str, Any]]] = {"yellow": [], "gamble": [], "fight": []}
    for cat in categories:
        chara_no_csv = ",".join(cat.codes)
        log(f"步骤：拉取近3个月明细（case/list）- {cat.title}", quiet=quiet)
        recent_rows = fetch_case_pages(
            opener,
            args.case_url,
            begin_dt=recent_begin_dt,
            end_dt=end_dt,
            new_ori_chara_codes_csv=chara_no_csv,
            timeout=args.timeout,
            page_size=args.page_size,
        )
        multi_addr_by_cat[cat.key] = build_multi_address_items(
            recent_rows=recent_rows,
            radius_m=20.0,
            min_count=2,
        )

    log("步骤：构建模板渲染数据（分地区/派出所）", quiet=quiet)
    basic_stats = compute_basic_stats(cases)
    fight_list = build_nested_list(cases, "打架斗殴")
    gamble_list = build_nested_list(cases, "赌博")
    sex_list = build_nested_list(cases, "涉黄")

    context = {
        "report_date": report_date_str,
        "srr_table_start_date": srr_start_dt.strftime("%Y-%m-%d"),
        "srr_table_end_date": end_dt.strftime("%Y-%m-%d"),
        "srr_rows_by_cat": srr_rows_by_cat,
        "total_count": basic_stats["total_count"],
        "type_summary_str": basic_stats["type_summary_str"],
        "stats_mock": stats_mock,
        "fight_list": fight_list,
        "gamble_list": gamble_list,
        "sex_list": sex_list,
        "repeat_fight": repeat_by_cat.get("fight", []),
        "repeat_yellow": repeat_by_cat.get("yellow", []),
        "repeat_gamble": repeat_by_cat.get("gamble", []),
        "repeat_min_count": int(args.repeat_min_count),
        "multi_addr_fight": multi_addr_by_cat.get("fight", []),
        "multi_addr_yellow": multi_addr_by_cat.get("yellow", []),
        "multi_addr_gamble": multi_addr_by_cat.get("gamble", []),
        "llm_analysis_problems": llm_probs,
        "llm_analysis_measures": llm_measures,
    }

    fmt = str(args.format or "html").strip().lower()
    if fmt not in ("html", "docx", "pdf"):
        raise ValueError(f"不支持的输出格式: {fmt}（仅支持 html/docx/pdf）")

    out_path = args.output.strip() or f"涉黄、赌、打架斗殴警情研判分析_{args.start_date}.{fmt}"
    if fmt == "html":
        log(f"步骤：渲染HTML模板 {args.template}", quiet=quiet)
        html_out = render_html(args.template, context)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(html_out)
    elif fmt == "docx":
        log("步骤：渲染DOCX", quiet=quiet)
        docx_bytes = render_report_docx(context)
        with open(out_path, "wb") as f:
            f.write(docx_bytes)
    else:
        log("步骤：渲染PDF", quiet=quiet)
        pdf_bytes = render_report_pdf(context)
        with open(out_path, "wb") as f:
            f.write(pdf_bytes)

    print(f"{fmt.upper()} 日报已生成: {out_path}")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("Interrupted", file=sys.stderr)
        raise
